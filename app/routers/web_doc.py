import os
import re
from io import BytesIO
from datetime import date, datetime

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from docx import Document
from docx.oxml.ns import qn

from app.core.database import get_db
from app.models.client import Client
from app.models.document_template import OfficeDocumentTemplate
from app.models.user import User
from app.services.document_templates import get_all_doc_specs, get_doc_title

router = APIRouter()

DOCS = get_all_doc_specs()


def _br_date(d: date | datetime | None) -> str:
    if not d:
        return ""
    if isinstance(d, datetime):
        d = d.date()
    return d.strftime("%d/%m/%Y")


def _safe(v) -> str:
    if v is None:
        return ""
    return str(v).strip()


def _sanitize_filename(name: str) -> str:
    name = (name or "").strip()
    name = re.sub(r'[\\/:*?"<>|]+', "-", name)
    name = re.sub(r"\s+", " ", name)
    return name


def _get_logged_user(request: Request, db: Session) -> User:
    user_id = request.session.get("user_id")
    office_id = request.session.get("office_id")

    if not user_id:
        raise HTTPException(status_code=401, detail="Usuário não autenticado")

    user = db.query(User).filter(User.id == user_id, User.is_active == True).first()
    if not user:
        raise HTTPException(status_code=401, detail="Usuário inválido")

    if not office_id:
        office_id = getattr(user, "office_id", None)

    if not office_id:
        raise HTTPException(status_code=403, detail="Escritório não identificado na sessão")

    return user


def _get_active_template_path(db: Session, office_id: int, doc_key: str) -> str:
    template = (
        db.query(OfficeDocumentTemplate)
        .filter(
            OfficeDocumentTemplate.office_id == office_id,
            OfficeDocumentTemplate.doc_key == doc_key,
            OfficeDocumentTemplate.is_active == True,
        )
        .order_by(OfficeDocumentTemplate.version.desc(), OfficeDocumentTemplate.id.desc())
        .first()
    )

    if not template:
        raise HTTPException(
            status_code=404,
            detail=(
                f"Não existe modelo ativo para '{get_doc_title(doc_key)}' "
                f"neste escritório. Cadastre o modelo em /doc-templates."
            ),
        )

    if not template.storage_path or not os.path.exists(template.storage_path):
        raise HTTPException(
            status_code=500,
            detail="O arquivo do modelo ativo não foi encontrado no armazenamento.",
        )

    return template.storage_path


# =========================================================
# Font helpers (Arial Narrow)
# =========================================================
def _apply_font(run, name: str = "Arial Narrow", size_pt: int | None = None) -> None:
    try:
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:ascii"), name)
        rfonts.set(qn("w:hAnsi"), name)
        rfonts.set(qn("w:cs"), name)
        rfonts.set(qn("w:eastAsia"), name)
    except Exception:
        try:
            run.font.name = name
        except Exception:
            pass

    if size_pt is not None:
        try:
            from docx.shared import Pt
            run.font.size = Pt(size_pt)
        except Exception:
            pass


def _ensure_paragraph_font(paragraph, font_name: str = "Arial Narrow") -> None:
    for r in paragraph.runs:
        _apply_font(r, font_name)


_PLACEHOLDER_NORMALIZER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


def _normalize_placeholders_in_text(text: str) -> str:
    def repl(match):
        inner = (match.group(1) or "").strip().lower()
        return "{{" + inner + "}}"

    return _PLACEHOLDER_NORMALIZER_RE.sub(repl, text or "")


# =========================================================
# 1) REPLACE PARA PLACEHOLDERS {{...}} (texto corrido)
# =========================================================
def _replace_placeholders_in_paragraph(
    paragraph,
    mapping: dict[str, str],
    font_name: str = "Arial Narrow",
) -> None:
    if not paragraph.runs:
        return

    full = "".join(run.text for run in paragraph.runs)
    full = _normalize_placeholders_in_text(full)

    if not full:
        return

    if "{{" not in full or "}}" not in full:
        _ensure_paragraph_font(paragraph, font_name)
        return

    nome_value = mapping.get("{{nome}}", "")

    has_nome = "{{nome}}" in full
    working = full.replace("{{nome}}", "__KJ_NOME__") if has_nome else full

    for k, v in mapping.items():
        if k == "{{nome}}":
            continue
        if k in working:
            working = working.replace(k, v or "")

    if (working == full) and (not has_nome):
        _ensure_paragraph_font(paragraph, font_name)
        return

    for r in paragraph.runs:
        r.text = ""

    if "__KJ_NOME__" not in working:
        paragraph.runs[0].text = working
        _apply_font(paragraph.runs[0], font_name)
        return

    before, after = working.split("__KJ_NOME__", 1)

    paragraph.runs[0].text = before
    _apply_font(paragraph.runs[0], font_name)

    run_nome = paragraph.add_run(nome_value or "")
    run_nome.bold = True
    _apply_font(run_nome, font_name)

    if after:
        run_after = paragraph.add_run(after)
        _apply_font(run_after, font_name)


def _replace_placeholders_everywhere(
    doc: Document,
    mapping: dict[str, str],
    font_name: str = "Arial Narrow",
) -> None:
    for p in doc.paragraphs:
        _replace_placeholders_in_paragraph(p, mapping, font_name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_placeholders_in_paragraph(p, mapping, font_name)

    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_placeholders_in_paragraph(p, mapping, font_name)
        for p in section.footer.paragraphs:
            _replace_placeholders_in_paragraph(p, mapping, font_name)


# =========================================================
# 2) FILL POR RÓTULO (compatibilidade com docs antigos)
# =========================================================
def _should_fill_label(full_text: str, label: str) -> bool:
    pos = full_text.find(label)
    if pos < 0:
        return False

    after = full_text[pos + len(label):]
    after_slice = after[:200]
    s = after_slice.strip()

    if not s:
        return True

    if re.fullmatch(r"[_\-\.\s]+", s):
        return True

    return False


def _replace_label_once(paragraph, label: str, value: str, font_name: str = "Arial Narrow") -> None:
    if not paragraph.runs:
        return

    full = "".join(run.text for run in paragraph.runs)
    if label not in full:
        _ensure_paragraph_font(paragraph, font_name)
        return

    if not _should_fill_label(full, label):
        _ensure_paragraph_font(paragraph, font_name)
        return

    pattern = re.escape(label) + r"\s*[_\-\.\s]*"
    replacement = f"{label} {value}".rstrip() if value else label
    full_final = re.sub(pattern, replacement + " ", full, count=1).rstrip()

    paragraph.runs[0].text = full_final
    _apply_font(paragraph.runs[0], font_name)

    for r in paragraph.runs[1:]:
        r.text = ""
        _apply_font(r, font_name)


def _fill_labels_everywhere(doc: Document, mapping_labels: dict[str, str], font_name: str = "Arial Narrow") -> None:
    for p in doc.paragraphs:
        for label, value in mapping_labels.items():
            _replace_label_once(p, label, value, font_name)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for label, value in mapping_labels.items():
                        _replace_label_once(p, label, value, font_name)

    for section in doc.sections:
        for p in section.header.paragraphs:
            for label, value in mapping_labels.items():
                _replace_label_once(p, label, value, font_name)

        for p in section.footer.paragraphs:
            for label, value in mapping_labels.items():
                _replace_label_once(p, label, value, font_name)


# =========================================================
# 3) MAPPINGS
# =========================================================
def _build_mapping_placeholders(c: Client) -> dict[str, str]:
    nome = _safe(c.nome)
    cpf = _safe(c.cpf_cnpj)
    rg = _safe(getattr(c, "rg", None))
    ssp_uf = _safe(getattr(c, "ssp_uf", None))
    est_civil = _safe(getattr(c, "estado_civil", None))
    profissao = _safe(getattr(c, "profissao", None))
    endereco = _safe(getattr(c, "endereco", None))
    telefone = _safe(getattr(c, "telefone", None))
    email = _safe(getattr(c, "email", None))
    nasc = _br_date(getattr(c, "nascimento", None))
    nacionalidade = _safe(getattr(c, "nacionalidade", None))

    return {
        "{{nome}}": nome,
        "{{nacionalidade}}": nacionalidade,
        "{{estado_civil}}": est_civil,
        "{{profissao}}": profissao,
        "{{rg}}": rg,
        "{{cpf}}": cpf,
        "{{ssp_uf}}": ssp_uf,
        "{{endereco}}": endereco,
        "{{telefone}}": telefone,
        "{{email}}": email,
        "{{nascimento}}": nasc,
    }


def _build_mapping_labels(c: Client) -> dict[str, str]:
    nome = _safe(c.nome)
    cpf = _safe(c.cpf_cnpj)
    rg = _safe(getattr(c, "rg", None))
    ssp_uf = _safe(getattr(c, "ssp_uf", None))
    est_civil = _safe(getattr(c, "estado_civil", None))
    profissao = _safe(getattr(c, "profissao", None))
    endereco = _safe(getattr(c, "endereco", None))
    telefone = _safe(getattr(c, "telefone", None))
    email = _safe(getattr(c, "email", None))
    nasc = _br_date(getattr(c, "nascimento", None))

    return {
        "Nome:": nome,
        "NOME:": nome,

        "CPF:": cpf,
        "CPF/CNPJ:": cpf,

        "RG N°:": rg,
        "RG Nº:": rg,
        "RG:": rg,

        "Órgão:": ssp_uf,
        "Orgao:": ssp_uf,
        "SSP-UF:": ssp_uf,

        "Estado Civil:": est_civil,
        "ESTADO CIVIL:": est_civil,

        "Profissão:": profissao,
        "Profissao:": profissao,

        "Data Nasc.:": nasc,
        "Data de Nascimento:": nasc,

        "Endereço Residencial:": endereco,
        "Endereco Residencial:": endereco,
        "Endereço:": endereco,

        "Telefone:": telefone,

        "Endereço eletrônico:": email,
        "Endereco eletronico:": email,
        "E-mail:": email,
        "Email:": email,
    }


@router.get("/docs/{doc_key}/{client_id}")
def gerar_documento(
    doc_key: str,
    client_id: int,
    request: Request,
    db: Session = Depends(get_db),
):
    if doc_key not in DOCS:
        raise HTTPException(status_code=404, detail="Documento inválido")

    user = _get_logged_user(request, db)
    office_id = request.session.get("office_id") or getattr(user, "office_id", None)

    if not office_id:
        raise HTTPException(status_code=403, detail="Escritório não identificado")

    cliente_query = db.query(Client).filter(Client.id == client_id)

    if hasattr(Client, "office_id"):
        cliente_query = cliente_query.filter(Client.office_id == office_id)

    cliente = cliente_query.first()

    if not cliente:
        raise HTTPException(status_code=404, detail="Cliente não encontrado")

    template_path = _get_active_template_path(db, office_id, doc_key)

    doc = Document(template_path)

    FONT_NAME = "Arial Narrow"

    _replace_placeholders_everywhere(doc, _build_mapping_placeholders(cliente), FONT_NAME)
    _fill_labels_everywhere(doc, _build_mapping_labels(cliente), FONT_NAME)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)

    titulo = get_doc_title(doc_key)
    nome_arquivo = _sanitize_filename(f"{titulo} - {cliente.nome}.docx")
    headers = {"Content-Disposition": f'attachment; filename="{nome_arquivo}"'}

    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )