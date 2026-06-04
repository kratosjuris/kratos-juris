import json
import os
import re
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document


# ---------------------------------------------------------------------------
# Specs de documentos
#
# IMPORTANTE:
#   - "memorial-calculo" foi REMOVIDO desta lista intencionalmente.
#     Ele é gerado 100% pelo sistema (sem upload de template).
#   - "cumprimento-voluntario" e "execucao-sentenca" aceitam papel timbrado
#     opcional. Se o escritório não tiver template cadastrado, o sistema
#     gera a petição automaticamente com texto jurídico embutido.
# ---------------------------------------------------------------------------

DOC_TEMPLATE_SPECS = {
    "procuracao": {
        "title": "Procuração",
        "allowed_placeholders": [
            "{{nome}}",
            "{{nacionalidade}}",
            "{{estado_civil}}",
            "{{profissao}}",
            "{{rg}}",
            "{{cpf}}",
            "{{ssp_uf}}",
            "{{endereco}}",
            "{{telefone}}",
            "{{email}}",
            "{{nascimento}}",
        ],
        "screen_text": (
            "OUTORGANTE: {{nome}}, {{estado_civil}}, {{profissao}}, "
            "portador(a) da cédula de identidade RG nº {{rg}}, inscrito(a) no CPF sob o nº {{cpf}}, "
            "residente e domiciliado(a) à {{endereco}}, telefone {{telefone}}, "
            "endereço eletrônico {{email}}."
        ),
    },
    "procuracao-a-rogo": {
        "title": "Procuração a rogo",
        "allowed_placeholders": [
            "{{nome}}",
            "{{nacionalidade}}",
            "{{estado_civil}}",
            "{{profissao}}",
            "{{rg}}",
            "{{cpf}}",
            "{{ssp_uf}}",
            "{{endereco}}",
            "{{telefone}}",
            "{{email}}",
            "{{nascimento}}",
        ],
        "screen_text": (
            "OUTORGANTE: {{nome}}, {{estado_civil}}, {{profissao}}, "
            "portador(a) da cédula de identidade RG nº {{rg}}, inscrito(a) no CPF sob o nº {{cpf}}, "
            "residente e domiciliado(a) à {{endereco}}, telefone {{telefone}}, "
            "endereço eletrônico {{email}}."
        ),
    },
    "hipossuficiencia": {
        "title": "Declaração de Hipossuficiência",
        "allowed_placeholders": [
            "{{nome}}",
            "{{nacionalidade}}",
            "{{estado_civil}}",
            "{{profissao}}",
            "{{rg}}",
            "{{cpf}}",
            "{{ssp_uf}}",
            "{{endereco}}",
            "{{telefone}}",
            "{{email}}",
            "{{nascimento}}",
        ],
        "screen_text": (
            "{{nome}}, {{estado_civil}}, {{profissao}}, "
            "portador(a) da cédula de identidade RG nº {{rg}}, inscrito(a) no CPF sob o nº {{cpf}}, "
            "residente e domiciliado(a) à {{endereco}}, telefone {{telefone}}, "
            "endereço eletrônico {{email}}."
        ),
    },
    "residencia": {
        "title": "Declaração de Residência",
        "allowed_placeholders": [
            "{{nome}}",
            "{{nacionalidade}}",
            "{{estado_civil}}",
            "{{profissao}}",
            "{{rg}}",
            "{{cpf}}",
            "{{ssp_uf}}",
            "{{endereco}}",
            "{{telefone}}",
            "{{email}}",
            "{{nascimento}}",
        ],
        "screen_text": (
            "Eu, {{nome}}, {{estado_civil}}, {{profissao}}, "
            "portador(a) da cédula de identidade RG nº {{rg}}, inscrito(a) no CPF sob o nº {{cpf}}, "
            "na falta de documento para comprovação de residência em nome próprio, DECLARO para os devidos fins, "
            "sob as penas da Lei, ser residente e domiciliado(a) à {{endereco}}, telefone {{telefone}}, "
            "endereço eletrônico {{email}}, Declaro ainda, estar ciente de que, se comprovadamente falsa a declaração, "
            "estar sujeito às sanções civis, administrativas e criminais previstas na legislação aplicável."
        ),
    },
    # memorial-calculo removido: gerado automaticamente pelo sistema
    "cumprimento-voluntario": {
        "title": "Papel Timbrado — Cumprimento Voluntário",
        "upload_hint": (
            "Faça upload do papel timbrado do escritório em .docx. "
            "Insira os placeholders abaixo onde desejar que os dados do cálculo apareçam. "
            "Se nenhum template for cadastrado, o sistema gerará a petição automaticamente."
        ),
        "allowed_placeholders": [
            "{{processo}}",
            "{{vara}}",
            "{{exequente}}",
            "{{executado}}",
            "{{data_calculo}}",
            "{{indice_correcao}}",
            "{{valor_original}}",
            "{{valor_corrigido}}",
            "{{juros}}",
            "{{multa}}",
            "{{honorarios}}",
            "{{custas}}",
            "{{deducoes}}",
            "{{total_bruto}}",
            "{{total_liquido}}",
            "{{memoria_calculo}}",
        ],
        "screen_text": (
            "Processo: {{processo}}. Vara: {{vara}}. "
            "Exequente: {{exequente}}. Executado: {{executado}}. "
            "Data do cálculo: {{data_calculo}}. "
            "Total líquido atualizado: {{total_liquido}}."
        ),
    },
    "execucao-sentenca": {
        "title": "Papel Timbrado — Execução de Sentença",
        "upload_hint": (
            "Faça upload do papel timbrado do escritório em .docx. "
            "Insira os placeholders abaixo onde desejar que os dados do cálculo apareçam. "
            "Se nenhum template for cadastrado, o sistema gerará a petição automaticamente."
        ),
        "allowed_placeholders": [
            "{{processo}}",
            "{{vara}}",
            "{{exequente}}",
            "{{executado}}",
            "{{data_calculo}}",
            "{{indice_correcao}}",
            "{{valor_original}}",
            "{{valor_corrigido}}",
            "{{juros}}",
            "{{multa}}",
            "{{honorarios}}",
            "{{custas}}",
            "{{deducoes}}",
            "{{total_bruto}}",
            "{{total_liquido}}",
            "{{memoria_calculo}}",
        ],
        "screen_text": (
            "Processo: {{processo}}. Vara: {{vara}}. "
            "Exequente: {{exequente}}. Executado: {{executado}}. "
            "Data do cálculo: {{data_calculo}}. "
            "Valor da execução: {{total_liquido}}."
        ),
    },
}


# =========================================================
# DIRETÓRIO BASE
# =========================================================
BASE_DIR = Path(__file__).resolve().parents[2]

DOC_TEMPLATES_STORAGE_DIR = Path(
    os.getenv(
        "DOC_TEMPLATES_STORAGE_DIR",
        str(BASE_DIR / "storage" / "offices"),
    )
).resolve()

DOC_TEMPLATES_MAX_MB    = int(os.getenv("DOC_TEMPLATES_MAX_MB", "5"))
DOC_TEMPLATES_MAX_BYTES = DOC_TEMPLATES_MAX_MB * 1024 * 1024

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


def get_all_doc_specs() -> dict:
    return DOC_TEMPLATE_SPECS


def get_doc_spec(doc_key: str) -> dict | None:
    return DOC_TEMPLATE_SPECS.get(doc_key)


def get_doc_title(doc_key: str) -> str:
    spec = get_doc_spec(doc_key)
    return spec["title"] if spec else doc_key


def get_allowed_placeholders(doc_key: str) -> list[str]:
    spec = get_doc_spec(doc_key)
    return spec["allowed_placeholders"] if spec else []


def sanitize_filename(name: str) -> str:
    name = (name or "").strip()
    name = re.sub(r'[\\/:*?"<>|]+', "-", name)
    name = re.sub(r"\s+", "_", name)
    return name


def ensure_storage_dir(office_id: int, doc_key: str) -> str:
    path = DOC_TEMPLATES_STORAGE_DIR / str(office_id) / "doc_templates" / doc_key
    path.mkdir(parents=True, exist_ok=True)
    return str(path)


def build_storage_path(
    office_id: int,
    doc_key: str,
    version: int,
    original_filename: str,
) -> str:
    base_dir  = Path(ensure_storage_dir(office_id, doc_key))
    safe_name = sanitize_filename(original_filename or f"{doc_key}.docx")
    timestamp = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
    return str((base_dir / f"v{version}_{timestamp}_{safe_name}").resolve())


def write_template_bytes(file_bytes: bytes, destination_path: str) -> None:
    destination = Path(destination_path).resolve()
    destination.parent.mkdir(parents=True, exist_ok=True)
    with open(destination, "wb") as f:
        f.write(file_bytes)


def remove_file_silently(path: str | None) -> None:
    try:
        if path:
            p = Path(path).resolve()
            if p.exists() and p.is_file():
                p.unlink()
    except Exception:
        pass


def _iter_doc_texts(doc: Document):
    for p in doc.paragraphs:
        yield p.text or ""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p.text or ""
    for section in doc.sections:
        for p in section.header.paragraphs:
            yield p.text or ""
        for p in section.footer.paragraphs:
            yield p.text or ""


def normalize_placeholder_name(name: str) -> str:
    return "{{" + (name or "").strip().lower() + "}}"


def extract_placeholders_from_docx_bytes(file_bytes: bytes) -> list[str]:
    doc = Document(BytesIO(file_bytes))
    placeholders = set()
    for text in _iter_doc_texts(doc):
        for match in _PLACEHOLDER_RE.finditer(text or ""):
            placeholders.add(normalize_placeholder_name(match.group(1)))
    return sorted(placeholders)


def validate_docx_placeholders(doc_key: str, file_bytes: bytes) -> dict:
    detected = extract_placeholders_from_docx_bytes(file_bytes)
    allowed  = set(get_allowed_placeholders(doc_key))
    invalid  = sorted([p for p in detected if p not in allowed])
    return {
        "doc_key":               doc_key,
        "detected_placeholders": detected,
        "invalid_placeholders":  invalid,
        "is_valid":              len(invalid) == 0,
    }


def detected_placeholders_to_json(placeholders: list[str]) -> str:
    return json.dumps(placeholders, ensure_ascii=False)