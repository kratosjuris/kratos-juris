"""
calculadora_documentos.py
=========================
Geração de documentos Word a partir dos resultados da calculadora jurídica.

Lógica por tipo de documento:
  - memorial-calculo       → gerado 100% pelo sistema (python-docx), sem template
  - cumprimento-voluntario → injeta dados no papel timbrado (template do escritório)
                             Se não houver template, gera automaticamente com texto jurídico
  - execucao-sentenca      → mesmo comportamento do cumprimento
"""

from __future__ import annotations

import os
import re
from io import BytesIO

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches
from fastapi import HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session

from app.models.document_template import OfficeDocumentTemplate
from app.services.calculos_utils import sanitize_filename


# ---------------------------------------------------------------------------
# Helpers gerais
# ---------------------------------------------------------------------------

def _apply_font(run, name="Arial", size_pt=None, bold=False, color=None):
    try:
        run.font.name = name
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), name)
        if size_pt:
            run.font.size = Pt(size_pt)
        if bold:
            run.font.bold = True
        if color:
            run.font.color.rgb = color
    except Exception:
        pass


_PLACEHOLDER_RE = re.compile(r"\{\{\s*([a-zA-Z0-9_]+)\s*\}\}")


def _normalize_placeholders(text):
    return _PLACEHOLDER_RE.sub(lambda m: "{{" + m.group(1).lower() + "}}", text or "")


def _replace_paragraph(paragraph, mapping):
    if not paragraph.runs:
        return
    full = _normalize_placeholders("".join(r.text for r in paragraph.runs))
    if "{{" not in full:
        for r in paragraph.runs:
            _apply_font(r)
        return
    for key, value in mapping.items():
        full = full.replace(key, value or "")
    paragraph.runs[0].text = full
    _apply_font(paragraph.runs[0])
    for r in paragraph.runs[1:]:
        r.text = ""


def _replace_everywhere(doc, mapping):
    for p in doc.paragraphs:
        _replace_paragraph(p, mapping)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    _replace_paragraph(p, mapping)
    for section in doc.sections:
        for p in section.header.paragraphs:
            _replace_paragraph(p, mapping)
        for p in section.footer.paragraphs:
            _replace_paragraph(p, mapping)


def _build_mapping(resultado):
    return {
        "{{processo}}":        resultado.get("processo", ""),
        "{{vara}}":            resultado.get("vara", ""),
        "{{exequente}}":       resultado.get("exequente", ""),
        "{{executado}}":       resultado.get("executado", ""),
        "{{data_calculo}}":    resultado.get("data_calculo", ""),
        "{{indice_correcao}}": resultado.get("indice_correcao", ""),
        "{{valor_original}}":  resultado.get("valor_original", ""),
        "{{valor_corrigido}}": resultado.get("valor_corrigido", ""),
        "{{juros}}":           resultado.get("juros", ""),
        "{{multa}}":           resultado.get("multa", ""),
        "{{honorarios}}":      resultado.get("honorarios", ""),
        "{{custas}}":          resultado.get("custas", ""),
        "{{deducoes}}":        resultado.get("total_deducoes", ""),
        "{{total_bruto}}":     resultado.get("total_bruto", ""),
        "{{total_liquido}}":   resultado.get("total_liquido", ""),
        "{{memoria_calculo}}": resultado.get("memoria_calculo", ""),
    }


def _get_active_template_path(db, office_id, doc_key):
    template = (
        db.query(OfficeDocumentTemplate)
        .filter(
            OfficeDocumentTemplate.office_id == office_id,
            OfficeDocumentTemplate.doc_key == doc_key,
            OfficeDocumentTemplate.is_active == True,
        )
        .order_by(
            OfficeDocumentTemplate.version.desc(),
            OfficeDocumentTemplate.id.desc(),
        )
        .first()
    )
    if not template:
        raise HTTPException(
            status_code=404,
            detail=f"Não existe modelo ativo para este documento: {doc_key}",
        )
    if not template.storage_path or not os.path.exists(template.storage_path):
        raise HTTPException(
            status_code=500,
            detail="O arquivo do modelo ativo não foi encontrado.",
        )
    return template.storage_path


def _streaming_response(bio, titulo, processo):
    bio.seek(0)
    processo_safe = sanitize_filename(processo or "sem-processo")
    filename = sanitize_filename(f"{titulo} - {processo_safe}.docx")
    return StreamingResponse(
        bio,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# ---------------------------------------------------------------------------
# Extração de valores de dano por tipo
# ---------------------------------------------------------------------------

def _extrair_danos(resultado: dict) -> dict:
    """
    Percorre os lançamentos e identifica valores de dano moral e material
    com base no campo 'tipo' de cada item.

    Retorna dict com:
      - dano_moral: str (valor formatado) ou None
      - dano_material: str (valor formatado) ou None
      - outros: lista de (tipo, valor_original)
    """
    itens = resultado.get("itens") or []

    dano_moral_val    = None
    dano_material_val = None
    outros = []

    for item in itens:
        tipo  = (item.get("tipo") or "").strip().upper()
        valor = item.get("valor_original", "")

        if "MORAL" in tipo:
            dano_moral_val = valor
        elif "MATERIAL" in tipo:
            dano_material_val = valor
        else:
            outros.append((item.get("tipo", ""), valor))

    return {
        "dano_moral":    dano_moral_val,
        "dano_material": dano_material_val,
        "outros":        outros,
    }


def _montar_frase_danos(danos: dict) -> str:
    """
    Monta a frase sobre os danos conforme o que foi informado.
    Suprime automaticamente o trecho que não tem valor.
    """
    partes = []

    if danos["dano_moral"]:
        partes.append(
            f"pelos danos morais no importe de {danos['dano_moral']}"
        )
    if danos["dano_material"]:
        partes.append(
            f"o valor de {danos['dano_material']} referente aos danos materiais"
        )

    if not partes:
        return ""

    return "Na sentença há a previsão da condenação " + " e ".join(partes) + "."


# ---------------------------------------------------------------------------
# Helpers de construção de parágrafo (compartilhado pelas petições)
# ---------------------------------------------------------------------------

def _par(doc, texto="", bold=False, size=12,
         align=WD_ALIGN_PARAGRAPH.JUSTIFY,
         space_before=0, space_after=6,
         color=None, indent_first=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(36)
    if texto:
        run = p.add_run(texto)
        _apply_font(run, "Arial", size, bold, color)
    return p


def _par_misto(doc, partes, size=12,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY,
               space_before=0, space_after=6,
               indent_first=False):
    """
    partes = lista de (texto, bold).
    Permite negrito inline dentro do mesmo parágrafo.
    """
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if indent_first:
        p.paragraph_format.first_line_indent = Pt(36)
    for texto, bold in partes:
        run = p.add_run(texto)
        _apply_font(run, "Arial", size, bold)
    return p


def _titulo_secao(doc, texto, size=12):
    p = _par(doc, texto, bold=True, size=size,
             align=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=12, space_after=6)
    return p


def _linha_assinatura(doc):
    _par(doc, "", space_before=24, space_after=4)
    _par(doc, "_" * 50, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _par(doc, "Advogado(a) — OAB/XX nº XXXXX", size=11,
         align=WD_ALIGN_PARAGRAPH.CENTER,
         color=RGBColor(0x44, 0x44, 0x44), space_after=0)


def _memoria_calculo_pagina(doc, memoria):
    if not memoria:
        return
    doc.add_page_break()
    _par(doc, "MEMÓRIA DE CÁLCULO", bold=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)
    for linha in memoria.splitlines():
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(linha)
        _apply_font(run, "Arial", 10, color=RGBColor(0x44, 0x44, 0x44))


def _cabecalho_peticao(doc, vara, processo, exequente, executado):
    """Cabeçalho padrão: endereçamento + partes."""
    # Endereçamento ao juízo — em negrito
    _par(doc,
         f"AO M.M JUÍZO DA {vara.upper()}",
         bold=True, size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER,
         space_after=12)

    # Processo em negrito — alinhado à esquerda
    _par_misto(doc,
               [("Processo nº ", False), (processo, True)],
               size=12,
               align=WD_ALIGN_PARAGRAPH.LEFT,
               space_after=4)

    # Partes — nomes em negrito e maiúsculas, alinhados à esquerda
    _par_misto(doc,
               [("Exequente: ", False), (exequente.upper(), True)],
               size=12,
               align=WD_ALIGN_PARAGRAPH.LEFT,
               space_after=2)
    _par_misto(doc,
               [("Executado: ", False), (executado.upper(), True)],
               size=12,
               align=WD_ALIGN_PARAGRAPH.LEFT,
               space_after=16)


# ---------------------------------------------------------------------------
# Petição de Cumprimento Voluntário
# ---------------------------------------------------------------------------

def _gerar_cumprimento_voluntario(resultado: dict, doc: Document = None) -> BytesIO:
    if doc is None:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Inches(1.18)
            section.bottom_margin = Inches(1.18)
            section.left_margin   = Inches(1.38)
            section.right_margin  = Inches(1.18)

    processo  = resultado.get("processo", "")
    vara      = resultado.get("vara", "")
    exequente = resultado.get("exequente", "")
    executado = resultado.get("executado", "")
    total     = resultado.get("total_liquido", "")
    data_calc = resultado.get("data_calculo", "")
    indice    = resultado.get("indice_correcao", "")
    memoria   = resultado.get("memoria_calculo", "")

    danos = _extrair_danos(resultado)
    frase_danos = _montar_frase_danos(danos)

    # ── Cabeçalho ──────────────────────────────────────────────────────────
    _cabecalho_peticao(doc, vara, processo, exequente, executado)

    # ── Qualificação ───────────────────────────────────────────────────────
    _par_misto(doc, [
        (exequente.upper(), True),
        (", já qualificado nos autos do processo em epígrafe, "
         "vem respeitosamente à presença de Vossa Excelência, por seu advogado "
         "que esta subscreve, com fundamento no art. 523 do Código de Processo "
         "Civil, apresentar o presente CUMPRIMENTO VOLUNTÁRIO DE SENTENÇA, "
         "expondo e requerendo o que segue:", False),
    ], size=12, space_before=6, space_after=12, indent_first=True)

    # ── Tópico I ───────────────────────────────────────────────────────────
    _titulo_secao(doc, "I – DO VALOR ATUALIZADO DA CONDENAÇÃO")

    _par(doc,
         "Considerando que nos autos da presente ação não cabe mais recurso, "
         "havendo, portanto, o trânsito em julgado, cabe ao Executado realizar "
         "o pagamento de forma voluntária no prazo de 15 (quinze) dias, "
         "consoante inteligência do art. 523, caput do CPC.",
         size=12, space_after=8, indent_first=True)

    # Frase de danos (suprimida se não houver)
    if frase_danos:
        _par(doc, frase_danos, size=12, space_after=8, indent_first=True)

    _par(doc,
         f"Desta forma, com base no cálculo elaborado em {data_calc}, "
         f"aplicando-se {indice}, o valor atualizado da condenação perfaz o "
         f"montante de {total} (por extenso), conforme memória de cálculo anexa.",
         size=12, space_after=8, indent_first=True)

    _par_misto(doc, [
        ("Portanto, cabe o executado ", False),
        (executado.upper(), True),
        (f" efetuar o pagamento do montante de {total} no prazo legal de "
         "15 (quinze) dias, conforme determina o art. 523, caput, do CPC, "
         "sob pena de multa de 10% (dez por cento) e honorários advocatícios "
         "de 10% (dez por cento) sobre o valor da condenação, nos termos do "
         "§ 1º do mesmo dispositivo legal.", False),
    ], size=12, space_after=16, indent_first=True)

    # ── Tópico III ─────────────────────────────────────────────────────────
    _titulo_secao(doc, "II – DOS REQUERIMENTOS")

    _par(doc, "Ante o exposto, requer:", size=12,
         space_after=4, indent_first=True)

    for letra, texto in [
        ("a)", f"A homologação do presente cumprimento voluntário de sentença, pelo valor de {total};"),
        ("b)", "A intimação do executado para efetuar o pagamento no prazo legal;"),
        ("c)", "Caso não haja o pagamento voluntário, seja expedido mandado de penhora e "
               "avaliação de bens, nos termos do art. 523, § 3º, do CPC."),
    ]:
        _par_misto(doc,
                   [(f"{letra} ", True), (texto, False)],
                   size=12, space_after=4, indent_first=True)

    _par(doc, "Nestes termos, pede deferimento.", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=0)

    _linha_assinatura(doc)
    _memoria_calculo_pagina(doc, memoria)

    bio = BytesIO()
    doc.save(bio)
    return bio


# ---------------------------------------------------------------------------
# Petição de Execução de Sentença
# ---------------------------------------------------------------------------

def _gerar_execucao_sentenca(resultado: dict, doc: Document = None) -> BytesIO:
    if doc is None:
        doc = Document()
        for section in doc.sections:
            section.top_margin    = Inches(1.18)
            section.bottom_margin = Inches(1.18)
            section.left_margin   = Inches(1.38)
            section.right_margin  = Inches(1.18)

    processo  = resultado.get("processo", "")
    vara      = resultado.get("vara", "")
    exequente = resultado.get("exequente", "")
    executado = resultado.get("executado", "")
    total     = resultado.get("total_liquido", "")
    data_calc = resultado.get("data_calculo", "")
    indice    = resultado.get("indice_correcao", "")
    memoria   = resultado.get("memoria_calculo", "")

    danos = _extrair_danos(resultado)
    frase_danos = _montar_frase_danos(danos)

    # Monta texto de soma dos danos para o tópico I
    if danos["dano_moral"] and danos["dano_material"]:
        frase_soma = (
            f"sendo {danos['dano_moral']} referente aos danos morais "
            f"e {danos['dano_material']} referente aos danos materiais"
        )
    elif danos["dano_moral"]:
        frase_soma = f"sendo {danos['dano_moral']} referente aos danos morais"
    elif danos["dano_material"]:
        frase_soma = f"sendo {danos['dano_material']} referente aos danos materiais"
    else:
        frase_soma = ""

    # ── Cabeçalho ──────────────────────────────────────────────────────────
    _cabecalho_peticao(doc, vara, processo, exequente, executado)

    # ── Qualificação ───────────────────────────────────────────────────────
    _par_misto(doc, [
        (exequente.upper(), True),
        (", já qualificado nos autos do processo em epígrafe, "
         "vem respeitosamente à presença de Vossa Excelência, por seu advogado "
         "que esta subscreve, com fundamento nos arts. 513 e seguintes do "
         "Código de Processo Civil, promover EXECUÇÃO DE SENTENÇA em face de ", False),
        (executado.upper(), True),
        (", pelos fundamentos e requerimentos a seguir:", False),
    ], size=12, space_before=6, space_after=12, indent_first=True)

    # ── Tópico I ───────────────────────────────────────────────────────────
    _titulo_secao(doc, "I – DOS FATOS")

    # Monta o parágrafo de fatos com os danos identificados
    if frase_soma:
        texto_fatos = (
            f"A sentença proferida nos presentes autos transitou em julgado, "
            f"condenando o executado ao pagamento das verbas apuradas, {frase_soma}. "
            "Decorrido o prazo para cumprimento voluntário sem o respectivo "
            "adimplemento, faz-se necessário o prosseguimento do feito com a "
            "presente execução."
        )
    else:
        texto_fatos = (
            "A sentença proferida nos presentes autos transitou em julgado, "
            "condenando o executado ao pagamento das verbas apuradas. "
            "Decorrido o prazo para cumprimento voluntário sem o respectivo "
            "adimplemento, faz-se necessário o prosseguimento do feito com a "
            "presente execução."
        )

    _par(doc, texto_fatos, size=12, space_after=12, indent_first=True)

    # ── Tópico II ──────────────────────────────────────────────────────────
    _titulo_secao(doc, "II – DO VALOR DA EXECUÇÃO")

    _par(doc,
         f"Conforme cálculo elaborado em {data_calc}, com aplicação de "
         f"{indice}, o valor atualizado da execução corresponde a {total} "
         "(por extenso), acrescido de multa de 10% (dez por cento) e "
         "honorários de 10% (dez por cento) sobre o valor da condenação, "
         "nos termos do art. 523, § 1º, do CPC, conforme memória de cálculo anexa.",
         size=12, space_after=12, indent_first=True)

    # ── Tópico III ─────────────────────────────────────────────────────────
    _titulo_secao(doc, "III – DOS REQUERIMENTOS")

    _par(doc, "Ante o exposto, requer:", size=12,
         space_after=4, indent_first=True)

    for letra, texto in [
        ("a)", f"O recebimento da presente execução pelo valor de {total};"),
        ("b)", "Seja realizada a penhora on-line pelo sistema SISBAJUD, "
               "nos termos do art. 854 do CPC;"),
        ("c)", "Caso não seja localizado valores na conta do Executado, requer "
               "desde já a expedição de mandado de penhora e avaliação de bens "
               "do executado, nos termos do art. 523, § 3º, do CPC;"),
        ("d)", "A intimação do executado, na pessoa de seu patrono, "
               "para ciência dos atos executivos."),
    ]:
        _par_misto(doc,
                   [(f"{letra} ", True), (texto, False)],
                   size=12, space_after=4, indent_first=True)

    _par(doc, "Nestes termos,", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_before=16, space_after=2)
    _par(doc, "Pede e espera deferimento.", size=12,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    _linha_assinatura(doc)
    _memoria_calculo_pagina(doc, memoria)

    bio = BytesIO()
    doc.save(bio)
    return bio


# ---------------------------------------------------------------------------
# Memorial de Cálculo — gerado 100% pelo sistema
# ---------------------------------------------------------------------------

def _gerar_memorial_calculo(resultado: dict) -> BytesIO:
    doc = Document()
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.18)
        section.right_margin  = Inches(1.18)

    CINZA = RGBColor(0x44, 0x44, 0x44)
    AZUL  = RGBColor(0x1F, 0x4E, 0x79)

    def linha_kv(label, value, size=10):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(3)
        r1 = p.add_run(f"{label}: ")
        _apply_font(r1, "Arial", size, bold=True, color=CINZA)
        r2 = p.add_run(value or "—")
        _apply_font(r2, "Arial", size, color=RGBColor(0, 0, 0))

    def titulo_secao(texto):
        p = _par(doc, texto, bold=True, size=11, align=WD_ALIGN_PARAGRAPH.LEFT,
                 color=AZUL, space_before=12, space_after=4)
        p.paragraph_format.keep_with_next = True
        pPr = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "1F4E79")
        pBdr.append(bottom)
        pPr.append(pBdr)
        return p

    # Cabeçalho
    _par(doc, "MEMORIAL DE CÁLCULO", bold=True, size=14,
         align=WD_ALIGN_PARAGRAPH.CENTER, color=AZUL, space_after=2)
    _par(doc, "Demonstrativo de atualização monetária conforme índices oficiais",
         size=9, align=WD_ALIGN_PARAGRAPH.CENTER,
         color=CINZA, space_after=12)

    titulo_secao("DADOS DO PROCESSO")
    linha_kv("Processo",      resultado.get("processo"))
    linha_kv("Vara",          resultado.get("vara"))
    linha_kv("Exequente",     resultado.get("exequente"))
    linha_kv("Executado",     resultado.get("executado"))
    linha_kv("Data do cálculo", resultado.get("data_calculo"))
    linha_kv("Índice aplicado", resultado.get("indice_correcao"))

    itens = resultado.get("itens") or []
    if itens:
        titulo_secao("VALORES ATUALIZADOS")
        from docx.enum.table import WD_TABLE_ALIGNMENT

        cols  = ["Data","Tipo","Descrição","Original","Fator","Corrigido","Juros desde","Juros","Total"]
        col_w = [1700, 1700, 2000, 1600, 1400, 1600, 1700, 1400, 1600]

        table = doc.add_table(rows=1, cols=len(cols))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdr = table.rows[0]
        for i, (cell, w) in enumerate(zip(hdr.cells, col_w)):
            cell.width = w
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(cols[i])
            _apply_font(run, "Arial", 9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            shd = OxmlElement("w:shd")
            shd.set(qn("w:val"), "clear")
            shd.set(qn("w:color"), "auto")
            shd.set(qn("w:fill"), "1F4E79")
            tcPr.append(shd)

        for idx, item in enumerate(itens):
            row = table.add_row()
            vals = [
                item.get("data",""), item.get("tipo",""), item.get("descricao",""),
                item.get("valor_original",""), item.get("fator",""),
                item.get("valor_corrigido",""), item.get("data_inicio_juros",""),
                item.get("juros",""), item.get("total",""),
            ]
            fill = "F2F2F2" if idx % 2 == 0 else "FFFFFF"
            for cell, val, w in zip(row.cells, vals, col_w):
                cell.width = w
                p = cell.paragraphs[0]
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(val)
                _apply_font(run, "Arial", 9)
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), fill)
                tcPr.append(shd)

    titulo_secao("RESUMO FINANCEIRO")
    for label, key in [
        ("Valor original",  "valor_original"),
        ("Valor corrigido", "valor_corrigido"),
        ("Juros",           "juros"),
        ("Multa",           "multa"),
        ("Honorários",      "honorarios"),
        ("Custas",          "custas"),
        ("Total bruto",     "total_bruto"),
        ("Deduções",        "total_deducoes"),
    ]:
        linha_kv(label, resultado.get(key))

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(6)
    r1 = p.add_run("TOTAL LÍQUIDO DEVIDO: ")
    _apply_font(r1, "Arial", 12, bold=True, color=AZUL)
    r2 = p.add_run(resultado.get("total_liquido", ""))
    _apply_font(r2, "Arial", 12, bold=True, color=AZUL)

    deducoes = resultado.get("deducoes") or []
    if deducoes:
        titulo_secao("DEDUÇÕES")
        for d in deducoes:
            linha_kv(d.get("nome", "Dedução"),
                     f"{d.get('tipo')} — {d.get('referencia')} = {d.get('valor_deduzido')}")

    memoria = resultado.get("memoria_calculo", "")
    if memoria:
        titulo_secao("MEMÓRIA DE CÁLCULO")
        for linha in memoria.splitlines():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(linha)
            _apply_font(run, "Arial", 9, color=CINZA)

    _par(doc, "", space_before=12)
    _par(doc,
         "Documento gerado eletronicamente pelo sistema de cálculos. "
         "Os índices de correção monetária são obtidos automaticamente "
         "das APIs oficiais do Banco Central do Brasil e do IBGE.",
         size=8, color=CINZA, align=WD_ALIGN_PARAGRAPH.CENTER)

    bio = BytesIO()
    doc.save(bio)
    return bio


# ---------------------------------------------------------------------------
# Petições via papel timbrado (template do escritório)
# ---------------------------------------------------------------------------

def _gerar_peticao_com_template(db, office_id, doc_key, resultado):
    """
    Abre o papel timbrado, limpa o body e gera a petição diretamente nele.
    Uma única passagem em memória — sem documento temporário intermediário.
    Header/footer com logo e endereço são preservados intactos.
    """
    template_path = _get_active_template_path(db, office_id, doc_key)
    doc_timbrado  = Document(template_path)

    W    = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    body = doc_timbrado.element.body

    # Preserva sectPr (margens, orientação, header/footer refs)
    sect_pr = body.find(f"{{{W}}}sectPr")

    # Limpa body mantendo apenas sectPr
    for child in list(body):
        tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
        if tag != "sectPr":
            body.remove(child)

    # Gera a petição diretamente no documento timbrado (sem doc temporário)
    if doc_key == "cumprimento-voluntario":
        _gerar_cumprimento_voluntario(resultado, doc=doc_timbrado)
    else:
        _gerar_execucao_sentenca(resultado, doc=doc_timbrado)

    # Os novos elementos foram adicionados ao body — move-os para antes do sectPr
    if sect_pr is not None:
        novos = [c for c in list(body)
                 if c.tag.split("}")[-1] not in ("sectPr",)
                 and list(body).index(c) > list(body).index(sect_pr)]
        # Na prática _par/_par_misto já appenda no body; só garantir ordem
        # (python-docx appenda antes do sectPr automaticamente via add_paragraph)
        pass

    bio = BytesIO()
    doc_timbrado.save(bio)
    return bio


# ---------------------------------------------------------------------------
# Ponto de entrada único
# ---------------------------------------------------------------------------

def gerar_docx_calculadora(
    db: Session,
    office_id: int,
    doc_key: str,
    resultado: dict,
) -> StreamingResponse:
    titulos = {
        "memorial-calculo":       "Memorial de Cálculo",
        "cumprimento-voluntario": "Cumprimento Voluntário",
        "execucao-sentenca":      "Execução de Sentença",
    }
    titulo   = titulos.get(doc_key, "Documento")
    processo = resultado.get("processo", "")

    if doc_key == "memorial-calculo":
        bio = _gerar_memorial_calculo(resultado)

    elif doc_key == "cumprimento-voluntario":
        try:
            bio = _gerar_peticao_com_template(db, office_id, doc_key, resultado)
        except HTTPException as e:
            if e.status_code == 404:
                bio = _gerar_cumprimento_voluntario(resultado)
            else:
                raise

    elif doc_key == "execucao-sentenca":
        try:
            bio = _gerar_peticao_com_template(db, office_id, doc_key, resultado)
        except HTTPException as e:
            if e.status_code == 404:
                bio = _gerar_execucao_sentenca(resultado)
            else:
                raise

    else:
        raise HTTPException(status_code=400, detail="Tipo de documento inválido.")

    return _streaming_response(bio, titulo, processo)